#!/usr/bin/env python3
"""
LLM Wiki Ingest — Properly handles .docx, .pdf, .txt, .md files.
Uses a two-step LLM approach:
  1. Analyze → get list of pages to create
  2. Generate → create each page individually

Usage:
    python3 ingest.py <file1> [file2] ...
    python3 ingest.py ~/Documents/report.pdf

Requires: pip3 install python-docx PyMuPDF
"""

import json
import sys
import re
import shutil
from pathlib import Path
from datetime import date
from urllib.request import Request, urlopen
from urllib.error import HTTPError

WIKI_DIR = Path(__file__).parent / "wiki"
WIKI_RAW = WIKI_DIR / "raw"
WIKI_PAGES = WIKI_DIR / "pages"
WIKI_INDEX = WIKI_DIR / "index.md"
WIKI_LOG = WIKI_DIR / "log.md"
AUTH_PATH = Path.home() / ".pi" / "agent" / "auth.json"

TODAY = date.today().isoformat()


# ─── Text Extraction ─────────────────────────────────────────────────────────

def extract_text_docx(filepath):
    from docx import Document
    doc = Document(str(filepath))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                try:
                    parts.append(f"{'#' * int(level)} {text}")
                except ValueError:
                    parts.append(f"## {text}")
            else:
                parts.append(text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def extract_text_pdf(filepath):
    import fitz
    doc = fitz.open(str(filepath))
    parts = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text().strip()
        if text:
            parts.append(f"[Page {page_num + 1}]\n{text}")
    doc.close()
    return "\n\n".join(parts)


def extract_text(filepath):
    suffix = filepath.suffix.lower()
    if suffix == ".docx":
        return extract_text_docx(filepath)
    elif suffix == ".pdf":
        return extract_text_pdf(filepath)
    elif suffix in (".md", ".txt", ".csv", ".tsv"):
        return filepath.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported: {suffix}")


# ─── LLM ─────────────────────────────────────────────────────────────────────

def load_token():
    with open(AUTH_PATH) as f:
        auth = json.load(f)
    token = auth.get("github-copilot", {}).get("access")
    if not token:
        print("ERROR: No Copilot token")
        sys.exit(1)
    return token


def call_copilot(system, prompt, token, model="claude-opus-4.6"):
    body = json.dumps({
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 4096,
        "temperature": 0.2,
        "stream": False
    }).encode()

    req = Request("https://api.githubcopilot.com/chat/completions", data=body, method="POST")
    req.add_header("Content-Type", "application/json")
    req.add_header("Authorization", f"Bearer {token}")
    req.add_header("Editor-Version", "vscode/1.99.0")
    req.add_header("Copilot-Integration-Id", "vscode-chat")
    req.add_header("Editor-Plugin-Version", "copilot-chat/0.25.0")

    resp = urlopen(req, timeout=120)
    data = json.loads(resp.read())
    return data["choices"][0]["message"]["content"]


# ─── Two-Step Ingest ──────────────────────────────────────────────────────────

def step1_analyze(filename, content, token):
    """Step 1: Analyze source and get list of pages to create"""
    # Take first 6k + last 2k for long docs
    if len(content) > 8000:
        content = content[:6000] + "\n\n[...middle omitted...]\n\n" + content[-2000:]

    prompt = f"""Analyze this source document and plan wiki pages to create.

Source: {filename}
Content:
---
{content}
---

Output a JSON array of pages to create. Each item:
{{
  "path": "sources/slug.md" or "entities/slug.md" or "concepts/slug.md",
  "title": "Human Readable Title",
  "type": "source" or "entity" or "concept",
  "summary": "One sentence describing what this page will contain"
}}

Rules:
- First item MUST be the source summary page (type: "source")
- Include ALL key entities: people, organizations, equipment, systems, locations, products
- Include ALL key concepts: methods, standards, regulations, processes, frameworks
- Aim for 5-15 pages total
- Slugs: lowercase-kebab-case

Output ONLY the JSON array, no markdown fences, no explanation."""

    response = call_copilot(
        "You are a wiki indexer. Output valid JSON only. No markdown code fences.",
        prompt, token
    )

    # Parse JSON array
    response = response.strip()
    if response.startswith("```"):
        response = re.sub(r'^```\w*\n?', '', response)
        response = re.sub(r'\n?```$', '', response)

    try:
        pages = json.loads(response)
        if isinstance(pages, list):
            return pages
    except json.JSONDecodeError:
        # Try to find array in response
        match = re.search(r'\[[\s\S]*\]', response)
        if match:
            try:
                return json.loads(match.group())
            except json.JSONDecodeError:
                pass

    print(f"  ⚠️  Could not parse page plan: {response[:200]}")
    return None


def step2_generate_page(page_info, filename, content, token):
    """Step 2: Generate a single wiki page"""
    # Shorter content for individual pages
    if len(content) > 6000:
        content = content[:4000] + "\n\n[...]\n\n" + content[-2000:]

    prompt = f"""Write a wiki page based on this source document.

Page to write:
- Path: {page_info['path']}
- Title: {page_info['title']}
- Type: {page_info['type']}
- Goal: {page_info['summary']}

Source: {filename}
Content:
---
{content}
---

Write the COMPLETE page in markdown with YAML frontmatter. Format:

---
title: "{page_info['title']}"
date: "{TODAY}"
tags: ["tag1", "tag2"]
source_count: 1
last_updated: "{TODAY}"
---

[page content here with [[Cross Reference]] links to other wiki pages]

Rules:
- Objective, third-person tone
- Use [[Page Title]] for cross-references to related concepts/entities
- Include all relevant facts from the source
- For source pages: provide a structured summary with key takeaways
- For entity pages: describe what it is, its role, relevant details
- For concept pages: define it, explain its application in context

Output ONLY the markdown page content (starting with ---). No JSON, no code fences."""

    return call_copilot(
        "You are a wiki writer. Output clean markdown with YAML frontmatter. No code fences around the output.",
        prompt, token
    )


# ─── Wiki Writing ─────────────────────────────────────────────────────────────

def write_page(rel_path, content):
    """Write a page to the wiki"""
    path = WIKI_PAGES / rel_path
    path.parent.mkdir(parents=True, exist_ok=True)
    # Clean up any code fences wrapping the whole content
    content = content.strip()
    if content.startswith("```markdown"):
        content = content[len("```markdown"):].strip()
    if content.startswith("```yaml"):
        content = content[len("```yaml"):].strip()
    if content.startswith("```"):
        content = content[3:].strip()
    if content.endswith("```"):
        content = content[:-3].strip()
    path.write_text(content, encoding="utf-8")


def update_index(pages_created):
    """Update index.md with new pages"""
    current = WIKI_INDEX.read_text(encoding="utf-8")

    sources = [p for p in pages_created if p["type"] == "source"]
    entities = [p for p in pages_created if p["type"] == "entity"]
    concepts = [p for p in pages_created if p["type"] == "concept"]

    additions = ""
    if sources:
        additions += "\n### Recent Sources\n"
        for p in sources:
            additions += f"- [[{p['title']}]] — {p['summary']}\n"
    if entities:
        additions += "\n### Entities\n"
        for p in entities:
            additions += f"- [[{p['title']}]] — {p['summary']}\n"
    if concepts:
        additions += "\n### Concepts\n"
        for p in concepts:
            additions += f"- [[{p['title']}]] — {p['summary']}\n"

    WIKI_INDEX.write_text(current + additions, encoding="utf-8")


def update_log(filename, num_pages):
    """Append to log"""
    entry = f"\n## [{TODAY}] ingest | {filename}\n\nProcessed via Copilot/claude-opus-4.6. Created {num_pages} pages.\n"
    current = WIKI_LOG.read_text(encoding="utf-8")
    WIKI_LOG.write_text(current + entry, encoding="utf-8")


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 ingest.py <file1.pdf> [file2.docx] ...")
        sys.exit(1)

    files = [Path(f) for f in sys.argv[1:]]
    token = load_token()

    print(f"🧠 LLM Wiki Ingest (two-step)")
    print(f"   Provider: GitHub Copilot (claude-opus-4.6)")
    print(f"   Files: {len(files)}")
    print(f"{'━' * 60}")

    for filepath in files:
        if not filepath.exists():
            print(f"\n❌ Not found: {filepath}")
            continue

        print(f"\n📄 {filepath.name} ({filepath.stat().st_size / 1024:.0f} KB)")

        # 1. Extract text
        print(f"   ↳ Extracting text...")
        try:
            text = extract_text(filepath)
        except Exception as e:
            print(f"   ❌ Extraction failed: {e}")
            continue

        if not text.strip():
            print(f"   ⚠️  No text extracted. Skipping.")
            continue

        print(f"   ✅ {len(text)} chars, {len(text.splitlines())} lines")

        # 2. Copy to wiki/raw/
        dest = WIKI_RAW / filepath.name
        if not dest.exists():
            shutil.copy2(filepath, dest)
            print(f"   ↳ Stored in wiki/raw/")

        # 3. Step 1: Analyze
        print(f"   ↳ Step 1: Analyzing document...")
        page_plan = step1_analyze(filepath.name, text, token)
        if not page_plan:
            continue

        print(f"   ✅ Plan: {len(page_plan)} pages to create")
        for p in page_plan:
            print(f"      • [{p['type']}] {p['title']}")

        # 4. Step 2: Generate each page
        pages_created = []
        for i, page_info in enumerate(page_plan):
            print(f"   ↳ Step 2: Writing [{i+1}/{len(page_plan)}] {page_info['path']}...")
            try:
                page_content = step2_generate_page(page_info, filepath.name, text, token)
                write_page(page_info["path"], page_content)
                pages_created.append(page_info)
                print(f"      ✅ Done")
            except HTTPError as e:
                err = e.read().decode()[:100]
                print(f"      ❌ HTTP {e.code}: {err}")
            except Exception as e:
                print(f"      ❌ Error: {e}")

        # 5. Update index and log
        if pages_created:
            update_index(pages_created)
            update_log(filepath.name, len(pages_created))
            print(f"   ✅ Complete: {len(pages_created)} pages written")

    # Summary
    print(f"\n{'━' * 60}")
    all_pages = sorted(WIKI_PAGES.rglob("*.md"))
    print(f"📊 Wiki: {len(all_pages)} total pages")
    for p in all_pages:
        print(f"   • {p.relative_to(WIKI_PAGES)}")


if __name__ == "__main__":
    main()
