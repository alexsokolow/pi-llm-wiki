#!/usr/bin/env python3
"""
Test ingesting real .docx and .pdf files from wiki/raw/ via GitHub Copilot.
Extracts text from binary formats, sends to LLM, and writes wiki pages.
"""

import json
import sys
import re
from pathlib import Path
from urllib.request import Request, urlopen
from urllib.error import HTTPError

WIKI_RAW = Path(__file__).parent / "wiki" / "raw"
WIKI_PAGES = Path(__file__).parent / "wiki" / "pages"
AUTH_PATH = Path.home() / ".pi" / "agent" / "auth.json"


def extract_text_docx(filepath: Path) -> str:
    """Extract text from .docx file"""
    from docx import Document
    doc = Document(str(filepath))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # Also get tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def extract_text_pdf(filepath: Path) -> str:
    """Extract text from .pdf file using PyMuPDF"""
    import fitz  # PyMuPDF
    doc = fitz.open(str(filepath))
    text_parts = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        if text.strip():
            text_parts.append(f"--- Page {page_num + 1} ---\n{text.strip()}")
    doc.close()
    return "\n\n".join(text_parts)


def extract_text(filepath: Path) -> str:
    """Extract text based on file extension"""
    suffix = filepath.suffix.lower()
    if suffix == ".docx":
        return extract_text_docx(filepath)
    elif suffix == ".pdf":
        return extract_text_pdf(filepath)
    elif suffix in (".md", ".txt", ".csv"):
        return filepath.read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def ingest_via_copilot(filename: str, content: str, token: str):
    """Send extracted text to Copilot for wiki processing"""
    # Load AGENTS.md as system prompt
    agents_md = (Path(__file__).parent / "AGENTS.md").read_text()

    # Truncate content if too long (leave room for prompt)
    max_content = 12000
    if len(content) > max_content:
        content = content[:max_content] + "\n\n[... truncated ...]"

    prompt = f"""Source document: {filename}

Content:
---
{content}
---

Process this according to the wiki schema. Output ONLY a JSON object with this exact shape:
{{
  "sourcePage": {{ "path": "sources/<slug>.md", "content": "full markdown with YAML frontmatter" }},
  "updates": [
    {{ "path": "entities/<slug>.md" OR "concepts/<slug>.md", "content": "full markdown with YAML frontmatter" }}
  ],
  "indexSnippet": "markdown lines to append to index.md"
}}

Rules:
- All content must be valid markdown with YAML frontmatter (title, date, tags, source_count, last_updated).
- Use [[Title]] cross-references between pages.
- The slug should be lowercase kebab-case.
- Extract key entities (people, orgs, systems, places) and concepts (methods, frameworks, standards).
- Be thorough — a source typically produces 5-15 wiki pages.
"""

    body = json.dumps({
        "model": "gpt-4o",
        "messages": [
            {"role": "system", "content": agents_md},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 4096,
        "stream": False
    }).encode()

    endpoint = "https://api.githubcopilot.com/chat/completions"
    req = Request(endpoint, data=body, method="POST")
    req.add_header("Content-Type", "application/json")
    req.add_header("Authorization", f"Bearer {token}")
    req.add_header("Editor-Version", "vscode/1.99.0")
    req.add_header("Copilot-Integration-Id", "vscode-chat")
    req.add_header("Editor-Plugin-Version", "copilot-chat/0.25.0")

    try:
        resp = urlopen(req, timeout=120)
        data = json.loads(resp.read())
        response_text = data["choices"][0]["message"]["content"]
        
        # Extract JSON from response
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            return json.loads(json_match.group())
        else:
            print(f"  ⚠️  No JSON found in response:")
            print(f"  {response_text[:500]}")
            return None
    except HTTPError as e:
        err_body = e.read().decode()
        print(f"  ❌ HTTP {e.code}: {err_body[:300]}")
        return None
    except json.JSONDecodeError as e:
        print(f"  ❌ JSON parse error: {e}")
        print(f"  Raw: {response_text[:300]}")
        return None
    except Exception as e:
        print(f"  ❌ Error: {e}")
        return None


def write_wiki_pages(result: dict):
    """Write the generated wiki pages to disk"""
    written = []
    
    # Write source page
    if result.get("sourcePage"):
        sp = result["sourcePage"]
        path = WIKI_PAGES / sp["path"]
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(sp["content"], encoding="utf-8")
        written.append(str(path.relative_to(WIKI_PAGES)))
    
    # Write entity/concept pages
    for upd in result.get("updates", []):
        path = WIKI_PAGES / upd["path"]
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(upd["content"], encoding="utf-8")
        written.append(str(path.relative_to(WIKI_PAGES)))
    
    # Append to index
    if result.get("indexSnippet"):
        index_path = Path(__file__).parent / "wiki" / "index.md"
        current = index_path.read_text(encoding="utf-8")
        index_path.write_text(current + "\n" + result["indexSnippet"] + "\n", encoding="utf-8")
        written.append("index.md (updated)")

    return written


def main():
    print("🧪 LLM Wiki — Binary File Ingest Test")
    print("=" * 50)
    
    # Load auth
    if not AUTH_PATH.exists():
        print(f"❌ No auth at {AUTH_PATH}")
        sys.exit(1)
    with open(AUTH_PATH) as f:
        auth = json.load(f)
    
    copilot = auth.get("github-copilot", {})
    token = copilot.get("access")
    if not token:
        print("❌ No Copilot access token")
        sys.exit(1)

    # List raw files
    raw_files = list(WIKI_RAW.iterdir())
    raw_files = [f for f in raw_files if not f.name.startswith(".")]
    
    if not raw_files:
        print("❌ No files in wiki/raw/")
        sys.exit(1)
    
    print(f"\n📁 Found {len(raw_files)} raw files:")
    for f in raw_files:
        print(f"   {f.name} ({f.stat().st_size / 1024:.0f} KB)")

    # Process each file
    for filepath in raw_files:
        print(f"\n{'━' * 50}")
        print(f"📄 Processing: {filepath.name}")
        
        # Step 1: Extract text
        print(f"  → Extracting text from {filepath.suffix}...")
        try:
            text = extract_text(filepath)
            print(f"  ✅ Extracted {len(text)} chars ({len(text.splitlines())} lines)")
            # Show first few lines
            preview = text[:300].replace("\n", "\n     ")
            print(f"     Preview: {preview}...")
        except Exception as e:
            print(f"  ❌ Extraction failed: {e}")
            continue
        
        # Step 2: Send to Copilot
        print(f"  → Sending to Copilot (gpt-4o)...")
        result = ingest_via_copilot(filepath.name, text, token)
        
        if not result:
            print(f"  ❌ Ingest failed for {filepath.name}")
            continue
        
        # Step 3: Write pages
        print(f"  → Writing wiki pages...")
        written = write_wiki_pages(result)
        print(f"  ✅ Created {len(written)} pages:")
        for w in written:
            print(f"     • {w}")

    # Final summary
    print(f"\n{'━' * 50}")
    print("📊 Final wiki state:")
    all_pages = list(WIKI_PAGES.rglob("*.md"))
    print(f"   Total pages: {len(all_pages)}")
    for p in sorted(all_pages):
        print(f"   • {p.relative_to(WIKI_PAGES)}")


if __name__ == "__main__":
    main()
